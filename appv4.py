from openai import OpenAI
import os
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import time
import sounddevice as sd
import numpy as np
import threading
from pydub import AudioSegment
from webdriver_manager.chrome import ChromeDriverManager
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# Initialize the OpenAI client with the API key from the environment variable
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


# Email sending functionality
def send_email(recipient_email: str, filename: str) -> None:
    """Sends an email with the meeting minutes DOCX file attached."""
    sender_email = "anasshaikh3168ss@gmail.com"
    sender_password = os.environ.get("GMAIL_APP_PASSWORD")  # Set this in your environment
    
    # Create message container
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = "Meeting Minutes Document"
    
    # Email body
    body = "Please find attached the meeting minutes document."
    msg.attach(MIMEText(body, 'plain'))
    
    # Attach DOCX file
    with open(filename, 'rb') as file:
        attach = MIMEApplication(file.read(), _subtype="docx")
        attach.add_header('Content-Disposition', 'attachment', filename=os.path.basename(filename))
        msg.attach(attach)
    
    # Send email via SMTP server
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, msg.as_string())
        print("\nEmail sent successfully!")
    except Exception as e:
        print(f"\nError sending email: {e}")
    finally:
        server.quit()

# Meeting Minutes Processing Functions
def transcribe_audio(audio_file_path: str) -> str:
    """Transcribes an audio file using the Whisper model."""
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcription.text

def abstract_summary_extraction(transcription: str) -> str:
    """Generates an abstract summary of the transcription."""
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a highly skilled AI trained in language comprehension and summarization. "
                    "Please read the following text and summarize it into a concise abstract paragraph. "
                    "Retain the most important points while avoiding unnecessary details."
                )
            },
            {"role": "user", "content": transcription}
        ]
    )
    summary = response.choices[0].message.content
    print("\n--- Abstract Summary ---")
    print(summary)
    return summary

def key_points_extraction(transcription: str) -> str:
    """Extracts the key points from the transcription."""
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a proficient AI that distills information into key points. "
                    "Based on the following text, list the main ideas, findings, or topics that capture the essence of the discussion."
                )
            },
            {"role": "user", "content": transcription}
        ]
    )
    key_points = response.choices[0].message.content
    print("\n--- Key Points ---")
    print(key_points)
    return key_points

def action_item_extraction(transcription: str) -> str:
    """Extracts action items from the transcription."""
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an AI expert in analyzing conversations and extracting actionable tasks. "
                    "Review the following text and identify any tasks, assignments, or actions mentioned. "
                    "List these action items clearly and concisely."
                )
            },
            {"role": "user", "content": transcription}
        ]
    )
    action_items = response.choices[0].message.content
    print("\n--- Action Items ---")
    print(action_items)
    return action_items

def sentiment_analysis(transcription: str) -> str:
    """Performs sentiment analysis on the transcription."""
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an expert in language and emotion analysis. "
                    "Analyze the sentiment of the following text by considering the overall tone, emotion conveyed, and context. "
                    "Indicate whether the sentiment is positive, negative, or neutral and provide brief explanations."
                )
            },
            {"role": "user", "content": transcription}
        ]
    )
    sentiment = response.choices[0].message.content
    print("\n--- Sentiment Analysis ---")
    print(sentiment)
    return sentiment

def meeting_minutes(transcription: str) -> dict:
    """Generates meeting minutes by combining the outputs of all analysis functions."""
    return {
        'abstract_summary': abstract_summary_extraction(transcription),
        'key_points': key_points_extraction(transcription),
        'action_items': action_item_extraction(transcription),
        'sentiment': sentiment_analysis(transcription)
    }

def save_as_docx(minutes: dict, filename: str) -> None:
    """Saves the meeting minutes to a Word document."""
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)
    print(f"\nMeeting minutes saved as {filename}")

# Selenium & Live Audio Recording Functions
chrome_options = Options()
chrome_options.add_argument("--disable-infobars")
chrome_options.add_argument("--disable-web-security")
chrome_options.add_argument("--disable-popup-blocking")
chrome_options.add_argument("--disable-blink-features=AutomationControlled")
chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
chrome_options.add_experimental_option("useAutomationExtension", False)
chrome_options.add_argument("--mute-audio")

def join_meeting(meeting_link: str, bot_name: str = "meet_reed.ai"):
    """Uses Selenium to join a Google Meet session."""
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    driver.get(meeting_link)
    time.sleep(5)
    
    try:
        name_input = driver.find_element(By.XPATH, "//input[@type='text']")
        name_input.clear()
        name_input.send_keys(bot_name + Keys.RETURN)
        print(f"Bot name set as '{bot_name}'")
    except Exception:
        print("No name input required, proceeding with meeting join...")
    
    print("Joined the meeting. Transcription will start after recording...")
    time.sleep(15)
    return driver

def record_audio_live(filename: str, duration: int = 120, sample_rate: int = 44100):
    """Records audio in real-time and saves it as a WAV file."""
    print("Recording audio in real-time...")
    audio_data = sd.rec(int(duration * sample_rate), samplerate=sample_rate, 
                       channels=1, dtype='int16', blocking=True)
    sd.wait()
    AudioSegment(
        np.array(audio_data, dtype=np.int16).tobytes(),
        frame_rate=sample_rate,
        sample_width=2,
        channels=1
    ).export(filename, format="wav")
    print(f"Audio saved to {filename}")

# Main Workflow
if __name__ == "__main__":
    # Get user email (in production, get this from Flask frontend)
    user_email = "anasshaikh1219@gmail.com"  # Replace with actual user email
    
    # Step 1: Join the Google Meet session
    meeting_link = "https://meet.google.com/pyc-voey-qpi"
    driver = join_meeting(meeting_link)
    
    # Ensure directories exist
    os.makedirs("audio", exist_ok=True)
    audio_file_path = "audio/meeting_audio_live03.wav"
    
    # Step 2: Record audio
    recording_thread = threading.Thread(target=record_audio_live, args=(audio_file_path, 120))
    recording_thread.start()
    time.sleep(90)
    recording_thread.join()
    
    # Step 3: Process audio and generate minutes
    transcription = transcribe_audio(audio_file_path)
    minutes = meeting_minutes(transcription)
    
    # Step 4: Save and email document
    doc_filename = "meeting_minutes03.docx"
    save_as_docx(minutes, doc_filename)
    send_email(user_email, doc_filename)
    
    # Cleanup
    driver.quit()
    print("\nMeeting exited and browser closed.")
    print("\nWork done Anas Sir")